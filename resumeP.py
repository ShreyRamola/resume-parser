import re
import pdfplumber
import docx
from tabulate import tabulate

class ResumeParser:
    def __init__(self, file_path):
        self.text = self._extract_text(file_path)
        self.file_path = file_path

    def _extract_text(self, file_path):
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_path.endswith('.pdf'):
            return self._extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please use .txt, .pdf, or .docx")

    def _extract_text_from_pdf(self, file_path):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    def _extract_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def extract_name(self):
        return next((line.strip() for line in self.text.splitlines()[:5] if line.strip()), "Name not found")

    def extract_contact_info(self):
        email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        phone_pattern = r'\b(?:\+?\d{1,3})?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        emails = re.findall(email_pattern, self.text)
        phones = re.findall(phone_pattern, self.text)
        return emails, phones

    def extract_skills(self):
        skills = ['Python', 'Java', 'SQL', 'C++', 'HTML', 'CSS', 'JavaScript', 'Excel', 'Linux']
        return [skill for skill in skills if skill.lower() in self.text.lower()]

    def parse_resume(self):
        name = self.extract_name()
        emails, phones = self.extract_contact_info()
        skills = self.extract_skills()
        return {
            "File": self.file_path,
            "Name": name,
            "Emails": ", ".join(emails) if emails else "Not found",
            "Phones": ", ".join(phones) if phones else "Not found",
            "Skills": ", ".join(skills) if skills else "Not found"
        }

if __name__ == "__main__":
    file_paths = input("Enter paths of resume files separated by commas (e.g., resume1.pdf, resume2.docx): ")
    files = [f.strip() for f in file_paths.split(",")]

    results = []
    for file_path in files:
        try:
            parser = ResumeParser(file_path)
            results.append(parser.parse_resume())
        except Exception as e:
            results.append({
                "File": file_path,
                "Name": "Error",
                "Emails": "Error",
                "Phones": "Error",
                "Skills": f"Error: {str(e)}"
            })

    print(tabulate(results, headers="keys", tablefmt="grid"))
import re
import pdfplumber
import docx
from tabulate import tabulate

class ResumeParser:
    def __init__(self, file_path):
        self.text = self._extract_text(file_path)
        self.file_path = file_path

    def _extract_text(self, file_path):
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_path.endswith('.pdf'):
            return self._extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please use .txt, .pdf, or .docx")

    def _extract_text_from_pdf(self, file_path):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    def _extract_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def extract_name(self):
        return next((line.strip() for line in self.text.splitlines()[:5] if line.strip()), "Name not found")

    def extract_contact_info(self):
        email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        phone_pattern = r'\b(?:\+?\d{1,3})?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        emails = re.findall(email_pattern, self.text)
        phones = re.findall(phone_pattern, self.text)
        return emails, phones

    def extract_skills(self):
        skills = ['Python', 'Java', 'SQL', 'C++', 'HTML', 'CSS', 'JavaScript', 'Excel', 'Linux']
        return [skill for skill in skills if skill.lower() in self.text.lower()]

    def parse_resume(self):
        name = self.extract_name()
        emails, phones = self.extract_contact_info()
        skills = self.extract_skills()
        return {
            "File": self.file_path,
            "Name": name,
            "Emails": ", ".join(emails) if emails else "Not found",
            "Phones": ", ".join(phones) if phones else "Not found",
            "Skills": ", ".join(skills) if skills else "Not found"
        }

if __name__ == "__main__":
    file_paths = input("Enter paths of resume files separated by commas (e.g., resume1.pdf, resume2.docx): ")
    files = [f.strip() for f in file_paths.split(",")]

    results = []
    for file_path in files:
        try:
            parser = ResumeParser(file_path)
            results.append(parser.parse_resume())
        except Exception as e:
            results.append({
                "File": file_path,
                "Name": "Error",
                "Emails": "Error",
                "Phones": "Error",
                "Skills": f"Error: {str(e)}"
            })

    print(tabulate(results, headers="keys", tablefmt="grid"))
